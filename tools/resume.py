"""RESUME MAKER — builds DOCX resumes from scratch.

Two types:
  - fresher   : ONE page. Dynamic qualifications list. Two visual variants.
  - ordinary  : Full-featured. Merges what was formerly "ordinary" + "detailed".
                Two visual variants. Every section has an enable/disable flag.

Education qualification options are stored in overrides.json under
  resume → edu_options   (list of strings)
and can be managed from the Dev panel on the resume page.
"""

import os
import json
import time
from io import BytesIO

from flask import Blueprint, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import OUTPUT_DIR, BASE_DIR
from tools import utils

bp = Blueprint("resume", __name__)

# ============================================================
#  Default education qualification options
#  (editable from Dev panel → resume → edu_options)
# ============================================================
DEFAULT_EDU_OPTIONS = [
    "SSC", "Diploma", "Inter", "Degree", "B.Tech", "M.Tech",
    "MBA", "MCA", "B.Com", "M.Com", "ITI", "Polytechnic"
]

OVERRIDES_PATH_RESUME = os.path.join(BASE_DIR, "overrides.json")


def _load_edu_options():
    try:
        with open(OVERRIDES_PATH_RESUME) as f:
            data = json.load(f)
        opts = data.get("resume", {}).get("edu_options")
        if opts and isinstance(opts, list):
            return opts
    except Exception:
        pass
    return DEFAULT_EDU_OPTIONS[:]


# ============================================================
#  Templates metadata (used by front-end picker)
# ============================================================
TEMPLATES = {
    "fresher_a": {
        "group": "fresher",
        "label": "Fresher — Classic",
        "blurb": "Clean single-page. Bold header, qualification list, personal profile.",
        "variant": "a",
    },
    "fresher_b": {
        "group": "fresher",
        "label": "Fresher — Compact",
        "blurb": "Minimal layout. Objective first, tabular qualifications, one-line profile.",
        "variant": "b",
    },
    "ordinary_a": {
        "group": "ordinary",
        "label": "Ordinary — Professional",
        "blurb": "Header + address, 5-col edu table, experience, strengths, personal details.",
        "variant": "a",
    },
    "ordinary_b": {
        "group": "ordinary",
        "label": "Ordinary — Detailed",
        "blurb": "Large name header, technical skills section, work experience, religion field.",
        "variant": "b",
    },
}

# ============================================================
#  Small docx helpers
# ============================================================
def _set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        sz = kwargs.get(side, 4)
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_margins(doc, top=0.7, bottom=0.7, left=0.7, right=0.7):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _para(doc_or_cell, text="", *, bold=False, size=11, align=None, italic=False,
          color=None, space_after=0):
    p = doc_or_cell.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def _section_heading(doc, text, *, size=12, underline=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.underline = underline
    return p


def _bullets(doc, items, size=11):
    for it in items:
        if not it:
            continue
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run("•  ")
        r1.font.size = Pt(size)
        r2 = p.add_run(str(it))
        r2.font.size = Pt(size)


def _kv_dotted(doc, label, value, *, size=11, label_w=2.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(label_w), WD_TAB_ALIGNMENT.LEFT)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run("\t: " + (value or ""))
    r2.font.size = Pt(size)
    return p


def _hdr_contact_table(doc, name, lines_left, lines_right, name_size=18):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    t.columns[0].width = Inches(4.0)
    t.columns[1].width = Inches(3.0)
    left, right = t.rows[0].cells
    pn = left.paragraphs[0]
    pn.paragraph_format.space_after = Pt(2)
    r = pn.add_run(name or "")
    r.bold = True
    r.font.size = Pt(name_size)
    for line in lines_left:
        if line:
            _para(left, line, size=10, space_after=0)
    for line in lines_right:
        if line:
            _para(right, line, size=10, align="right", space_after=0)
    for cell in (left, right):
        _set_cell_border(cell, top=0, left=0, right=0, bottom=0)
    return t


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _edu_table_5col(doc, edu, headers):
    edu = [e for e in edu if any((e.get(k) or "").strip()
                                 for k in ("c1", "c2", "c3", "c4", "c5"))]
    if not edu:
        _para(doc, "(No education rows added.)", size=11, italic=True)
        return
    t = doc.add_table(rows=1 + len(edu), cols=5)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        _shade_cell(c, "DCE6F1")
        _set_cell_border(c, top=6, bottom=6, left=6, right=6)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(edu):
        cells = t.rows[ri + 1].cells
        for ci, key in enumerate(["c1", "c2", "c3", "c4", "c5"]):
            c = cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run((row.get(key) or "").strip())
            r.font.size = Pt(10)
            _set_cell_border(c, top=4, bottom=4, left=4, right=4)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ============================================================
#  FRESHER BUILDERS
# ============================================================
def build_fresher_a(d):
    """Classic fresher — same proven layout, one page."""
    doc = Document()
    _set_margins(doc, top=0.5, bottom=0.5, left=0.6, right=0.6)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    sec = d.get("enabled_sections", {})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESUME")
    r.bold = True; r.underline = True; r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    name = (d.get("name") or "").strip()
    father = (d.get("father") or "").strip()
    _para(doc, name, bold=True, size=13)
    if father:
        _para(doc, f"D/O  {father}", size=11)

    for line in [d.get("addr1"), d.get("addr2"), d.get("city")]:
        line = (line or "").strip()
        if line:
            _para(doc, line, size=10, space_after=0)
    if (d.get("mobile") or "").strip():
        _para(doc, f"Mobile No : {d['mobile'].strip()}", size=10, space_after=0)
    if (d.get("email") or "").strip():
        _para(doc, f"Email Id  : {d['email'].strip()}", size=10)
    _hr(doc)

    if sec.get("objective", True):
        _section_heading(doc, "CAREER OBJECTIVE:")
        obj = (d.get("objective") or "").strip() or (
            "To obtain a challenging and responsible position in an organisation "
            "that contributes towards its growth using my abilities and knowledge.")
        _para(doc, obj, size=11, align="justify")

    if sec.get("qualifications", True):
        _section_heading(doc, "ACADEMIC QUALIFICATIONS:")
        quals = [q for q in (d.get("qualifications") or [])
                 if (q.get("course") or "").strip() or (q.get("institute") or "").strip()]
        if not quals:
            _para(doc, "(No qualifications added.)", size=11, italic=True)
        else:
            for q in quals:
                course = (q.get("course") or "").strip()
                inst   = (q.get("institute") or "").strip()
                year   = (q.get("year") or "").strip()
                line = course
                if inst:  line += "   —   " + inst
                if year:  line += f"   ({year})"
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(1)
                p.add_run(line).font.size = Pt(11)

    if sec.get("experience", True):
        _section_heading(doc, "EXPERIENCE:")
        _para(doc, "Fresher.", size=11)

    if sec.get("profile", True):
        _section_heading(doc, "PERSONAL PROFILE:")
        for label, key in [
            ("Date of Birth", "dob"), ("Gender", "gender"),
            ("Nationality", "nationality"), ("Marital Status", "marital"),
            ("Languages Known", "languages"), ("Hobbies", "hobbies"),
        ]:
            v = (d.get(key) or "").strip()
            if v:
                _kv_dotted(doc, label, v, size=11, label_w=1.7)

    if sec.get("declaration", True):
        _section_heading(doc, "DECLARATION:")
        _para(doc, "I hereby declare that the information furnished above is true to the best of my knowledge.",
              size=11, align="justify")
        place = (d.get("place") or "").strip()
        date  = (d.get("date") or "").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(f"Place : {place}").font.size = Pt(11)
        p.add_run("\tSignature").font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.add_run(f"Date  : {date}").font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(8)
        p3.add_run(f"({name})").font.size = Pt(11)

    return doc


def build_fresher_b(d):
    """Compact fresher — slightly different header style, tabular qualifications."""
    doc = Document()
    _set_margins(doc, top=0.5, bottom=0.5, left=0.65, right=0.65)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    sec = d.get("enabled_sections", {})
    name = (d.get("name") or "").strip()

    # Large centered name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True; r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    # Contact line centered
    contact_parts = []
    if (d.get("mobile") or "").strip():  contact_parts.append(d["mobile"].strip())
    if (d.get("email") or "").strip():   contact_parts.append(d["email"].strip())
    if contact_parts:
        _para(doc, "  |  ".join(contact_parts), size=10, align="center")

    # Address line
    addr = [x for x in [d.get("addr1"), d.get("addr2"), d.get("city")] if (x or "").strip()]
    if addr:
        _para(doc, ", ".join(a.strip() for a in addr), size=10, align="center", space_after=4)
    _hr(doc)

    father = (d.get("father") or "").strip()
    if father:
        _para(doc, f"D/O  {father}", size=11)

    if sec.get("objective", True):
        _section_heading(doc, "CAREER OBJECTIVE:")
        obj = (d.get("objective") or "").strip() or (
            "Seeking a challenging position to apply my skills and contribute "
            "to organisational growth while furthering my professional development.")
        _para(doc, obj, size=11, align="justify")

    if sec.get("qualifications", True):
        _section_heading(doc, "ACADEMIC QUALIFICATIONS:")
        quals = [q for q in (d.get("qualifications") or [])
                 if (q.get("course") or "").strip() or (q.get("institute") or "").strip()]
        if quals:
            t = doc.add_table(rows=1 + len(quals), cols=3)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t.autofit = True
            for i, h in enumerate(["Qualification", "Institution", "Year"]):
                c = t.rows[0].cells[i]
                c.text = ""
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p.add_run(h); r2.bold = True; r2.font.size = Pt(10)
                _shade_cell(c, "DCE6F1")
                _set_cell_border(c, top=6, bottom=6, left=6, right=6)
            for ri, q in enumerate(quals):
                cells = t.rows[ri + 1].cells
                for ci, key in enumerate(["course", "institute", "year"]):
                    c = cells[ci]; c.text = ""
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run((q.get(key) or "").strip()).font.size = Pt(10)
                    _set_cell_border(c, top=4, bottom=4, left=4, right=4)

    if sec.get("experience", True):
        _section_heading(doc, "EXPERIENCE:")
        _para(doc, "Fresher.", size=11)

    if sec.get("profile", True):
        _section_heading(doc, "PERSONAL PROFILE:")
        for label, key in [
            ("Date of Birth", "dob"), ("Gender", "gender"),
            ("Nationality", "nationality"), ("Marital Status", "marital"),
            ("Languages Known", "languages"), ("Hobbies", "hobbies"),
        ]:
            v = (d.get(key) or "").strip()
            if v:
                _kv_dotted(doc, label, v, size=11, label_w=1.7)

    if sec.get("declaration", True):
        _section_heading(doc, "DECLARATION:")
        _para(doc, "I hereby declare that the information furnished above is true to the best of my knowledge.",
              size=11, align="justify")
        place = (d.get("place") or "").strip()
        date  = (d.get("date") or "").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(f"Place : {place}").font.size = Pt(11)
        p.add_run("\tSignature").font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.add_run(f"Date  : {date}").font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(8)
        p3.add_run(f"({name})").font.size = Pt(11)

    return doc


# ============================================================
#  ORDINARY BUILDERS
# ============================================================
def build_ordinary_a(d):
    """Ordinary — Professional layout (address header + 5-col edu + experience)."""
    doc = Document()
    _set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    sec = d.get("enabled_sections", {})
    name = (d.get("name") or "").strip()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESUME")
    r.bold = True; r.underline = True; r.font.size = Pt(15)

    addr_lines = [d.get("addr1"), d.get("addr2"), d.get("city")]
    contact_lines = []
    if (d.get("mobile") or "").strip():
        contact_lines.append("Mobile : " + d["mobile"].strip())
    if (d.get("email") or "").strip():
        contact_lines.append("Mail Id : " + d["email"].strip())
    _hdr_contact_table(doc, name, addr_lines, contact_lines, name_size=15)
    _hr(doc)

    if sec.get("objective", True):
        _section_heading(doc, "CAREER OBJECTIVE")
        obj = (d.get("objective") or "").strip() or (
            "Looking for a growth-oriented organisation to contribute my services "
            "thereby developing as an effective professional.")
        _para(doc, obj, size=11, align="justify")

    if sec.get("education", True):
        _section_heading(doc, "EDUCATIONAL QUALIFICATION")
        headers = ["ACADEMIC\nQUALIFICATION", "NAME OF INSTITUTE",
                   "BOARD / UNIVERSITY", "YEAR OF PASSING", "PERCENTAGE"]
        _edu_table_5col(doc, d.get("education") or [], headers)

    if sec.get("experience", True):
        work = [x for x in (d.get("experience") or []) if x.strip()]
        _section_heading(doc, "EXPERIENCE")
        if work:
            _bullets(doc, work)
        else:
            _para(doc, "Fresher.", size=11)

    if sec.get("strengths", True):
        strengths = [s for s in (d.get("strengths") or []) if s.strip()]
        if strengths:
            _section_heading(doc, "STRENGTHS")
            _bullets(doc, strengths)

    if sec.get("profile", True):
        _section_heading(doc, "PERSONAL PROFILE")
        for label, key in [
            ("Name", "name"), ("Father's Name", "father"),
            ("Date of Birth", "dob"), ("Gender", "gender"),
            ("Nationality", "nationality"), ("Marital Status", "marital"),
            ("Languages Known", "languages"), ("Hobbies", "hobbies"),
        ]:
            v = (d.get(key) or "").strip()
            if v:
                _kv_dotted(doc, label, v, size=11, label_w=1.7)

    if sec.get("declaration", True):
        _section_heading(doc, "DECLARATION")
        _para(doc, "I hereby declare that the above information is true to the best of my knowledge.",
              size=11, align="justify")
        place = (d.get("place") or "").strip()
        date  = (d.get("date") or "").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(f"Place : {place}").font.size = Pt(11)
        p.add_run("\tSignature").font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.add_run(f"Date  : {date}").font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(8)
        p3.add_run(f"({name})").font.size = Pt(11)

    return doc


def build_ordinary_b(d):
    """Ordinary — Detailed layout (large name header + technical skills + religion)."""
    doc = Document()
    _set_margins(doc, top=0.7, bottom=0.7, left=0.8, right=0.8)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    sec = d.get("enabled_sections", {})
    name = (d.get("name") or "").strip()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESUME")
    r.bold = True; r.underline = True; r.font.size = Pt(15)
    p.paragraph_format.space_after = Pt(4)

    contact = []
    if (d.get("email") or "").strip():  contact.append("Email: " + d["email"].strip())
    if (d.get("mobile") or "").strip(): contact.append("Cell: " + d["mobile"].strip())
    _hdr_contact_table(doc, name, [], contact, name_size=14)

    if sec.get("objective", True):
        _section_heading(doc, "CAREER OBJECTIVE:")
        obj = (d.get("objective") or "").strip() or (
            "To obtain a challenging and responsible position in an organisation, "
            "contributing towards its growth using my abilities and knowledge.")
        _para(doc, obj, size=11, align="justify")

    if sec.get("education", True):
        _section_heading(doc, "EDUCATIONAL QUALIFICATION:")
        headers = ["Qualification", "University / College Name", "Specialization",
                   "Percentage / Marks", "Year of Passing"]
        edu = d.get("education") or []
        edu_f = [e for e in edu if any((e.get(k) or "").strip()
                                       for k in ("c1","c2","c3","c4","c5"))]
        if not edu_f:
            _para(doc, "(No education rows added.)", size=11, italic=True)
        else:
            widths = [Inches(1.1), Inches(2.0), Inches(1.3), Inches(1.0), Inches(1.1)]
            t = doc.add_table(rows=1 + len(edu_f), cols=5)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for col, w in zip(t.columns, widths):
                col.width = w
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = ""; c.width = widths[i]
                p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(h); r2.bold = True; r2.font.size = Pt(10)
                _shade_cell(c, "DCE6F1"); _set_cell_border(c, top=6, bottom=6, left=6, right=6)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for ri, row in enumerate(edu_f):
                cells = t.rows[ri + 1].cells
                for ci, key in enumerate(["c1","c2","c3","c4","c5"]):
                    c = cells[ci]; c.width = widths[ci]; c.text = ""
                    p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.add_run((row.get(key) or "").strip()).font.size = Pt(10)
                    _set_cell_border(c, top=4, bottom=4, left=4, right=4)
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if sec.get("experience", True):
        work = [x for x in (d.get("experience") or []) if x.strip()]
        _section_heading(doc, "WORK EXPERIENCE:")
        if work:
            _bullets(doc, work)
        else:
            _para(doc, "Fresher.", size=11)

    if sec.get("skills", True):
        skills = [s for s in (d.get("skills") or [])
                  if (s.get("label") or "").strip() or (s.get("value") or "").strip()]
        if skills:
            _section_heading(doc, "TECHNICAL SKILLS:")
            for s in skills:
                label = (s.get("label") or "").strip()
                value = (s.get("value") or "").strip()
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.35)
                p2.paragraph_format.space_after = Pt(2)
                tabs = p2.paragraph_format.tab_stops
                tabs.add_tab_stop(Inches(2.6), WD_TAB_ALIGNMENT.LEFT)
                p2.add_run("•  ").font.size = Pt(11)
                p2.add_run(label).font.size = Pt(11)
                p2.add_run(f"\t: {value}").font.size = Pt(11)

    if sec.get("strengths", True):
        strengths = [s for s in (d.get("strengths") or []) if s.strip()]
        if strengths:
            _section_heading(doc, "STRENGTHS:")
            for s in strengths:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.35)
                p2.paragraph_format.space_after = Pt(2)
                p2.add_run("➢  " + s).font.size = Pt(11)

    if sec.get("profile", True):
        _section_heading(doc, "PERSONAL DETAILS:")
        profile_fields = [
            ("Name", "name"), ("Father's Name", "father"),
            ("Date of Birth", "dob"), ("Gender", "gender"),
            ("Religion", "religion"), ("Marital Status", "marital"),
            ("Nationality", "nationality"), ("Languages Known", "languages"),
            ("Hobbies", "hobbies"),
        ]
        for label, key in profile_fields:
            v = (d.get(key) or "").strip()
            if v:
                _kv_dotted(doc, label, v, size=11, label_w=1.9)

        addr = [(d.get(k) or "").strip() for k in ("addr1","addr2","city")]
        addr = [a for a in addr if a]
        if addr:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            r1 = p2.add_run("Communication Address"); r1.bold = True; r1.font.size = Pt(11)
            tabs = p2.paragraph_format.tab_stops
            tabs.add_tab_stop(Inches(1.9), WD_TAB_ALIGNMENT.LEFT)
            r2 = p2.add_run("\t: " + addr[0]); r2.font.size = Pt(11)
            for line in addr[1:]:
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_after = Pt(0)
                p3.paragraph_format.left_indent = Inches(2.05)
                p3.add_run(line).font.size = Pt(11)

    if sec.get("declaration", True):
        _section_heading(doc, "DECLARATION:")
        _para(doc, "I hereby declare that the above information is true to the best of my knowledge.",
              size=11, align="justify")
        place = (d.get("place") or "").strip()
        date  = (d.get("date") or "").strip()
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(8)
        tabs = p2.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
        p2.add_run(f"Place : {place}").font.size = Pt(11)
        p2.add_run("\tSignature").font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.add_run(f"Date  : {date}").font.size = Pt(11)
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p4.paragraph_format.space_before = Pt(8)
        p4.add_run(f"({name})").font.size = Pt(11)

    return doc


BUILDERS = {
    "fresher_a":  build_fresher_a,
    "fresher_b":  build_fresher_b,
    "ordinary_a": build_ordinary_a,
    "ordinary_b": build_ordinary_b,
}

# ============================================================
#  Routes
# ============================================================
@bp.route("/resume")
def page():
    edu_options = _load_edu_options()
    return render_template("tool_resume.html",
                           templates=TEMPLATES,
                           edu_options=edu_options)


@bp.route("/resume/edu-options", methods=["GET"])
def get_edu_options():
    return jsonify({"options": _load_edu_options()})


@bp.route("/resume/edu-options", methods=["POST"])
def set_edu_options():
    j = request.get_json(silent=True) or {}
    opts = j.get("options")
    if not isinstance(opts, list):
        return jsonify({"error": "options must be a list"})
    opts = [str(o).strip() for o in opts if str(o).strip()]
    try:
        with open(OVERRIDES_PATH_RESUME) as f:
            data = json.load(f)
    except Exception:
        data = {}
    if "resume" not in data:
        data["resume"] = {}
    data["resume"]["edu_options"] = opts
    with open(OVERRIDES_PATH_RESUME, "w") as f:
        json.dump(data, f, indent=2)
    return jsonify({"ok": True, "options": opts})


@bp.route("/resume/build", methods=["POST"])
def build():
    j = request.get_json(silent=True) or {}
    tpl_key = j.get("template")
    if tpl_key not in BUILDERS:
        return jsonify({"error": f"Unknown template: {tpl_key}"})

    fields = j.get("fields") or {}
    payload = dict(fields)
    payload["education"]         = j.get("education") or []
    payload["experience"]        = j.get("experience") or []
    payload["strengths"]         = j.get("strengths") or []
    payload["skills"]            = j.get("skills") or []
    payload["qualifications"]    = j.get("qualifications") or []
    payload["enabled_sections"]  = j.get("enabled_sections") or {}

    try:
        doc = BUILDERS[tpl_key](payload)
    except Exception as e:
        return jsonify({"error": f"Build failed: {e}"})

    safe_name = (fields.get("name") or "resume").strip().replace(" ", "_")
    safe_name = "".join(c for c in safe_name if c.isalnum() or c in "-_") or "resume"
    variant_label = TEMPLATES.get(tpl_key, {}).get("label", tpl_key).replace(" ", "_").replace("—", "-")
    out_name = f"{safe_name}_{variant_label}_{int(time.time())}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return jsonify({"out": f"/file/{out_name}", "name": out_name})
