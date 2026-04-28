"""CONVERT — file format conversions.

Supports the conversions you actually use:
  - PDF              → JPG (one image per page)
  - PDF              → DOCX  (uses pdf2docx)
  - Images           → single PDF (merge), or separate PDFs
  - Image (any fmt)  → JPG / PNG / WEBP
  - Image            → DOCX (centered on A4)
  - HEIC / HEIF / AVIF / WEBP / PNG / GIF / BMP all read as input.

Outputs are written to OUTPUT_DIR and downloaded via /file/<name>.
"""

import os
import io
import time
import tempfile
import zipfile
from flask import Blueprint, render_template, request, jsonify, send_file

import fitz                                 # PyMuPDF
from PIL import Image
import pillow_heif

from docx import Document
from docx.shared import Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import utils
from config import OUTPUT_DIR

pillow_heif.register_heif_opener()

bp = Blueprint("convert", __name__)

IMG_EXTS = {"jpg","jpeg","png","webp","heic","heif","avif","bmp","gif","tif","tiff"}


@bp.route("/convert")
def page():
    return render_template("tool_convert.html", title="Convert")


# ---------- helpers ----------
def _flatten(img):
    """RGBA / palette → solid white background, return RGB."""
    if img.mode in ("RGBA","LA") or (img.mode == "P" and "transparency" in img.info):
        img = img.convert("RGBA")
        bg = Image.new("RGB", img.size, (255,255,255))
        bg.paste(img, mask=img.split()[-1])
        return bg
    return img.convert("RGB")


def _save_uploads(files):
    saved = []
    for f in files:
        if not f or not f.filename:
            continue
        uid, path = utils.save_upload(f)
        saved.append({"uid": uid, "path": path, "name": f.filename,
                      "ext": utils.file_ext(f.filename)})
    return saved


def _outname(uid, name, new_ext):
    base = os.path.splitext(name)[0]
    safe = "".join(c if c.isalnum() or c in "-_." else "_" for c in base)
    return f"{uid}_{safe}.{new_ext}"


def _outpath(name):
    return os.path.join(OUTPUT_DIR, name)


def _public(name):
    return f"/file/{name}"


# ---------- ROUTE: convert ----------
@bp.route("/convert/run", methods=["POST"])
def run():
    op = request.form.get("op", "img2jpg")
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "Pick at least one file"})
    saved = _save_uploads(files)
    outs = []

    try:
        if op == "pdf2jpg":
            for f in saved:
                if f["ext"] != "pdf":
                    outs.append({"src": f["name"], "error": "not a PDF"}); continue
                try:
                    doc = fitz.open(f["path"])
                    for i in range(len(doc)):
                        pix = doc.load_page(i).get_pixmap(matrix=fitz.Matrix(2, 2))
                        stem = os.path.splitext(f["name"])[0]
                        safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in stem)
                        on = f"{f['uid']}_{safe}_p{i+1}.jpg"
                        pix.save(_outpath(on))
                        outs.append({"src": f["name"], "out": _public(on),
                                     "label": f"page {i+1}"})
                    doc.close()
                except Exception as e:
                    outs.append({"src": f["name"], "error": str(e)})

        elif op == "pdf2docx":
            from pdf2docx import Converter
            for f in saved:
                if f["ext"] != "pdf":
                    outs.append({"src": f["name"], "error": "not a PDF"}); continue
                on = _outname(f["uid"], f["name"], "docx")
                try:
                    cv = Converter(f["path"])
                    cv.convert(_outpath(on))
                    cv.close()
                    outs.append({"src": f["name"], "out": _public(on), "label": "docx"})
                except Exception as e:
                    outs.append({"src": f["name"], "error": str(e)})

        elif op in ("img2jpg", "img2png", "img2webp"):
            target = {"img2jpg":"jpg","img2png":"png","img2webp":"webp"}[op]
            for f in saved:
                if f["ext"] not in IMG_EXTS:
                    outs.append({"src": f["name"], "error": "not an image"}); continue
                try:
                    with Image.open(f["path"]) as im:
                        rgb = _flatten(im)
                        on = _outname(f["uid"], f["name"], target)
                        if target == "jpg":
                            rgb.save(_outpath(on), "JPEG", quality=92)
                        elif target == "png":
                            rgb.save(_outpath(on), "PNG")
                        else:
                            rgb.save(_outpath(on), "WEBP", quality=92)
                    outs.append({"src": f["name"], "out": _public(on), "label": target})
                except Exception as e:
                    outs.append({"src": f["name"], "error": str(e)})

        elif op in ("imgs2pdf_one", "imgs2pdf_each"):
            imgs = []
            names = []
            for f in saved:
                if f["ext"] not in IMG_EXTS: continue
                try:
                    im = Image.open(f["path"])
                    imgs.append(_flatten(im).copy())
                    names.append(f["name"])
                    im.close()
                except Exception:
                    pass
            if not imgs:
                return jsonify({"error":"No valid images"})
            if op == "imgs2pdf_one":
                merged = f"{int(time.time())}_merged.pdf"
                imgs[0].save(_outpath(merged), save_all=True, append_images=imgs[1:])
                outs.append({"src":"merged", "out": _public(merged), "label":"merged.pdf"})
            else:
                for im, n in zip(imgs, names):
                    on = _outname(saved[0]["uid"], n, "pdf")
                    im.save(_outpath(on))
                    outs.append({"src": n, "out": _public(on), "label":"pdf"})

        elif op == "img2docx":
            for f in saved:
                if f["ext"] not in IMG_EXTS:
                    outs.append({"src": f["name"], "error": "not an image"}); continue
                doc = Document()
                sec = doc.sections[0]
                sec.page_height = Mm(297); sec.page_width = Mm(210)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # Convert to JPG temp first so docx accepts it (no HEIC etc.)
                tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
                with Image.open(f["path"]) as im:
                    _flatten(im).save(tmp.name, "JPEG", quality=92)
                run.add_picture(tmp.name, width=Inches(6.0))
                on = _outname(f["uid"], f["name"], "docx")
                doc.save(_outpath(on))
                os.unlink(tmp.name)
                outs.append({"src": f["name"], "out": _public(on), "label":"docx"})

        else:
            return jsonify({"error": f"Unknown op '{op}'"})

    except Exception as e:
        return jsonify({"error": f"{type(e).__name__}: {e}"})

    return jsonify({"results": outs})


# ---------- multi-file ZIP download ----------
@bp.route("/convert/zip", methods=["POST"])
def zip_results():
    j = request.get_json(silent=True) or {}
    paths = j.get("paths") or []
    if not paths:
        return jsonify({"error":"no paths"}), 400
    zname = f"convert_{int(time.time())}.zip"
    zpath = _outpath(zname)
    with zipfile.ZipFile(zpath, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in paths:
            name = os.path.basename(p)
            full = _outpath(name)
            if os.path.exists(full):
                zf.write(full, name)
    return jsonify({"out": _public(zname)})
